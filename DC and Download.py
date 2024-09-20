import os
import streamlit as st
import bcrypt
import pandas as pd
from langchain_openai import OpenAI
from langchain.globals import set_verbose
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import tempfile

# Set verbosity (True to enable debug information, False to disable it)
set_verbose(False)

# Set up OpenAI API key
os.environ['OPENAI_API_KEY'] = "Enter API Key"

# Initialize the OpenAI LLM with the specified temperature
llm = OpenAI(temperature=0.6)

# File path for storing user credentials
USER_DATA_FILE = "users.csv"

# Function to read and extract text from a .docx file
def read_docx(file):
    doc = Document(file)
    text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text.strip())
    return text

# Function to split text into smaller chunks
def split_text(text, max_tokens=1000):
    sentences = text.split('\n')
    chunks = []
    current_chunk = []
    current_length = 0
    for sentence in sentences:
        sentence_length = len(sentence.split())
        if current_length + sentence_length > max_tokens:
            chunks.append("\n".join(current_chunk))
            current_chunk = [sentence]
            current_length = sentence_length
        else:
            current_chunk.append(sentence)
            current_length += sentence_length
    if current_chunk:
        chunks.append("\n".join(current_chunk))
    return chunks

# Function to compare key points and suggest missing ones
def compare_key_points(master_text, client_text):
    if master_text == client_text:
        return "Both documents are identical. No differences found."
    
    master_chunks = split_text(master_text)
    client_chunks = split_text(client_text)
    missing_key_points = []

    for master_chunk in master_chunks:
        for client_chunk in client_chunks:
            if len(missing_key_points) >= 20:
                break  # Stop if we already have 20 key points
            prompt = f"""
            The following text is from a master policy document:

            {master_chunk}

            The following text is from a client policy document:

            {client_chunk}

            Identify the key points that are present in the master policy document but missing in the client policy document. Focus only on the differences where the master document contains essential clauses or sections that are not present in the client document.
            Limit the output to the top 20 key points.
            """
            response = llm.invoke(prompt)
            for point in response.split('\n'):
                if point.strip() and point not in client_text:
                    missing_key_points.append((point, client_chunk))

    return missing_key_points[:20]  # Limit to the top 20 key points

# Function to provide suggestions for updating the client document
def generate_suggestions(master_text, client_text):
    master_chunks = split_text(master_text)
    client_chunks = split_text(client_text)
    suggestions = []

    for i, (master_chunk, client_chunk) in enumerate(zip(master_chunks, client_chunks)):
        if len(suggestions) >= 20:
            break  # Stop if we already have 20 suggestions
        prompt = f"""
        The following text is from a master policy document:

        {master_chunk}

        The following text is from a client policy document:

        {client_chunk}

        Compare the two documents and provide detailed suggestions for updating the client policy document to better align with the master policy document. Focus on missing clauses, differences in wording, and contextual deviations. Only suggest changes to the client documents. Do not suggest changes to the master document. Focus only on identifying any differences between the two documents. Limit the suggestions to the top 20 points and highlight the key points clearly.
        """
        response = llm.invoke(prompt)
        suggestions.extend(response.split('\n'))  # Collect suggestions

    return "\n\n".join(suggestions[:20])  # Limit to top 20 suggestions

# Function to insert missing sections into the client document at the appropriate location
def insert_missing_sections(missing_key_points, client_doc):
    doc = Document(client_doc)
    for point, client_chunk in missing_key_points:
        inserted = False
        for para in doc.paragraphs:
            if client_chunk.strip() in para.text and not inserted:
                # Find the paragraph where the chunk is present and insert the point after it
                insert_index = doc.paragraphs.index(para) + 1
                insert_para = doc.add_paragraph(point)
                insert_para.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
                inserted = True
                break
        if not inserted:
            # If no match is found, it means the entire chunk is missing, so we insert it in the closest possible location
            closest_para = None
            min_distance = float('inf')
            for i, para in enumerate(doc.paragraphs):
                if point.strip() not in para.text:
                    # Calculate the "distance" to find the closest insertion point
                    distance = abs(len(para.text) - len(client_chunk))
                    if distance < min_distance:
                        min_distance = distance
                        closest_para = i

            if closest_para is not None:
                insert_para = doc.add_paragraph(point)
                insert_para.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
                doc.paragraphs.insert(closest_para + 1, insert_para)
            else:
                # If no suitable location is found, append at the end
                insert_para = doc.add_paragraph(point)
                insert_para.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

    return doc

# Utility functions for user authentication
def load_user_data():
    if os.path.exists(USER_DATA_FILE):
        return pd.read_csv(USER_DATA_FILE)
    else:
        return pd.DataFrame(columns=["username", "password_hash"])

def save_user_data(users_df):
    users_df.to_csv(USER_DATA_FILE, index=False)

def hash_password(password):
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

def check_password(password, hashed):
    return bcrypt.checkpw(password.encode(), hashed.encode())

def authenticate(username, password):
    users_df = load_user_data()
    if username in users_df['username'].values:
        stored_hashed_password = users_df[users_df['username'] == username]['password_hash'].values[0]
        return check_password(password, stored_hashed_password)
    return False

def signup(username, password):
    users_df = load_user_data()
    if username in users_df['username'].values:
        return False  # Username already exists
    hashed_password = hash_password(password)
    new_user_df = pd.DataFrame({"username": [username], "password_hash": [hashed_password]})
    users_df = pd.concat([users_df, new_user_df], ignore_index=True)
    save_user_data(users_df)
    return True

# Page navigation
def main():
    if 'page' not in st.session_state:
        st.session_state.page = 'signup'

    if st.session_state.page == 'signup':
        signup_page()
    elif st.session_state.page == 'login':
        login_page()
    elif st.session_state.page == 'upload':
        upload_page()
    elif st.session_state.page == 'results':
        results_page()

def signup_page():
    st.title("🔐 Signup Page")
    username = st.text_input("Choose a Username")
    password = st.text_input("Choose a Password", type="password")
    signup_button = st.button("Signup")

    if signup_button:
        if signup(username, password):
            st.success("Signup successful! Please log in.")
            st.session_state.page = 'login'
            st.rerun()
        else:
            st.error("Username already exists. Please choose a different username.")

    st.write("Already have an account? [Login](#)", unsafe_allow_html=True)
    if st.button("Go to Login"):
        st.session_state.page = 'login'
        st.rerun()

def login_page():
    st.title("🔐 Login Page")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    login_button = st.button("Login")

    if login_button:
        if authenticate(username, password):
            st.success("Login successful!")
            st.session_state.page = 'upload'
            st.rerun()
        else:
            st.error("Invalid credentials. Please try again.")

    st.write("Don't have an account? [Signup](#)", unsafe_allow_html=True)
    if st.button("Go to Signup"):
        st.session_state.page = 'signup'
        st.rerun()

def upload_page():
    st.title("📄 Upload Policy Documents")
    st.sidebar.button("Logout", on_click=lambda: logout())

    master_policy_file = st.file_uploader("📁 Upload Master Policy Document", type=["docx"])
    client_policy_files = st.file_uploader("📂 Upload Client Policy Documents to Compare", type=["docx"], accept_multiple_files=True)

    if st.button("Compare Documents"):
        if master_policy_file is None:
            st.warning("Please upload the master policy document.")
        elif not client_policy_files:
            st.warning("Please upload at least one client policy document.")
        else:
            st.session_state.master_policy_text = "\n".join(read_docx(master_policy_file))
            st.session_state.client_policy_files = client_policy_files
            st.session_state.page = 'results'
            st.rerun()

def results_page():
    st.title("🔍 Comparison Results")
    st.sidebar.button("Logout", on_click=lambda: logout())
    st.sidebar.button("Back to Upload Page", on_click=lambda: back_to_upload())

    master_policy_text = st.session_state.master_policy_text
    client_policy_files = st.session_state.client_policy_files

    for client_file in client_policy_files:
        client_policy_text = "\n".join(read_docx(client_file))
        st.success(f"✅ Client Policy Document '{client_file.name}' loaded successfully.")

        # Compare key points and identify missing ones
        st.subheader(f"📊 Comparing '{client_file.name}' with Master Document")
        missing_key_points = compare_key_points(master_policy_text, client_policy_text)

        with st.expander(f"🔎 Missing Key Points in '{client_file.name}'", expanded=True):
            for point, _ in missing_key_points:
                st.write(f"- {point}")

        # Provide suggestions for improvements
        suggestions = generate_suggestions(master_policy_text, client_policy_text)
        with st.expander(f"💡 Suggestions for '{client_file.name}'", expanded=False):
            st.write(suggestions)

        # Insert missing sections into the client document
        updated_client_doc = insert_missing_sections(missing_key_points, client_file)
        
        # Save the updated document using tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            updated_client_doc.save(tmp.name)
            tmp.seek(0)
            st.download_button(
                label=f"⬇️ Download Updated '{client_file.name}'",
                data=tmp.read(),
                file_name=f"updated_{client_file.name}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

def logout():
    st.session_state.page = 'login'
    st.session_state.master_policy_text = None
    st.session_state.client_policy_files = None
    st.rerun()

def back_to_upload():
    st.session_state.page = 'upload'
    st.rerun()

if __name__ == '__main__':
    main()
